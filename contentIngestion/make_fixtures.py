"""Generate a real .docx and .pdf fixture for end-to-end ingest testing."""
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

doc = Document()
doc.add_heading("Machine Learning", level=1)
doc.add_paragraph("Intro to the course.")
doc.add_heading("Gradient Descent", level=2)
doc.add_paragraph("Gradient descent minimizes a loss function by following "
                  "the negative gradient of the loss with respect to weights.")
doc.add_heading("Backpropagation", level=2)
doc.add_paragraph("Backprop computes gradients through the network using the "
                  "chain rule, layer by layer from the output back to the input.")
doc.save("/tmp/lecture.docx")

c = canvas.Canvas("/tmp/lecture.pdf", pagesize=letter)
text = c.beginText(72, 720)
for line in ["Machine Learning",
             "Gradient descent minimizes a loss function by following the",
             "negative gradient. Backpropagation computes gradients through",
             "the network using the chain rule, layer by layer."]:
    text.textLine(line)
c.drawText(text)
c.showPage()
c.save()
print("wrote /tmp/lecture.docx and /tmp/lecture.pdf")

"""DOCX extractor: one unit per heading section, tracking the heading path.

Mirrors the markdown extractor's shape — Word "Heading N" paragraph styles drive
the section boundaries and heading_path; body paragraphs accumulate under them.
"""
from __future__ import annotations
import io
import re
from typing import List, Tuple

from backend.models import ChunkPosition
from backend.chunking import ExtractedUnit

_HEADING_STYLE = re.compile(r"heading\s+(\d+)", re.IGNORECASE)


class DocxExtractor:
    """python-docx paragraph walk; sections split on Heading styles."""

    def extract(self, data: bytes) -> List[ExtractedUnit]:
        from docx import Document
        doc = Document(io.BytesIO(data))
        sections = self._split_sections(doc)
        if len(sections) > 1:
            return [ExtractedUnit(text=b, position=ChunkPosition(heading_path=p))
                    for p, b in sections]
        return self._paragraph_units(doc)

    def _split_sections(self, doc) -> List[Tuple[List[str], str]]:
        """Accumulate body paragraphs under the active heading path."""
        sections: List[Tuple[List[str], str]] = []
        path: List[str] = []
        body: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            level = self._heading_level(para)
            if level:
                self._flush(sections, path, body)
                path = path[: level - 1] + [text]
                body = []
            else:
                body.append(text)
        self._flush(sections, path, body)
        return sections

    def _heading_level(self, para) -> int:
        """Return the heading level from the paragraph style, or 0 if body."""
        name = getattr(para.style, "name", "") or ""
        if name.strip().lower() == "title":
            return 1
        m = _HEADING_STYLE.match(name.strip())
        return int(m.group(1)) if m else 0

    def _flush(self, sections, path, body) -> None:
        """Append the accumulated body as a section when non-empty."""
        text = "\n".join(body).strip()
        if text:
            sections.append((list(path), text))

    def _paragraph_units(self, doc) -> List[ExtractedUnit]:
        """No headings: emit each non-empty paragraph as an unstructured unit."""
        return [ExtractedUnit(text=p.text.strip(), structured=False)
                for p in doc.paragraphs if p.text.strip()]
